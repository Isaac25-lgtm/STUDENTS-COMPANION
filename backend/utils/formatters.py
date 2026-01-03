"""
Output Formatting Module
Formats results for frontend display and export
"""
import json
import pandas as pd
import numpy as np
from typing import Dict, Any, List
from datetime import datetime


class ResultFormatter:
    """
    Formats statistical results for various outputs
    """
    
    @staticmethod
    def to_json_safe(obj: Any) -> Any:
        """Convert object to JSON-serializable format"""
        if isinstance(obj, pd.DataFrame):
            return obj.to_dict('records')
        elif isinstance(obj, pd.Series):
            return obj.to_dict()
        elif isinstance(obj, np.ndarray):
            return obj.tolist()
        elif isinstance(obj, (np.integer, np.floating)):
            return float(obj)
        elif isinstance(obj, np.bool_):
            return bool(obj)
        elif pd.isna(obj):
            return None
        elif isinstance(obj, dict):
            return {k: ResultFormatter.to_json_safe(v) for k, v in obj.items()}
        elif isinstance(obj, (list, tuple)):
            return [ResultFormatter.to_json_safe(item) for item in obj]
        else:
            return obj
    
    @staticmethod
    def format_results_response(results: Dict, table: Dict = None, interpretation: Dict = None) -> Dict[str, Any]:
        """
        Format complete results response for frontend
        """
        response = {
            'timestamp': datetime.now().isoformat(),
            'statistical_results': ResultFormatter.to_json_safe(results)
        }
        
        if table:
            response['apa_table'] = ResultFormatter.to_json_safe(table)
        
        if interpretation:
            response['ai_interpretation'] = interpretation
        
        return response
    
    @staticmethod
    def format_error_response(error: str, details: Dict = None) -> Dict[str, Any]:
        """Format error response"""
        return {
            'success': False,
            'error': error,
            'details': details,
            'timestamp': datetime.now().isoformat()
        }
    
    @staticmethod
    def format_success_response(data: Any, message: str = None) -> Dict[str, Any]:
        """Format success response"""
        return {
            'success': True,
            'data': ResultFormatter.to_json_safe(data),
            'message': message,
            'timestamp': datetime.now().isoformat()
        }


class WordDocumentGenerator:
    """
    Generate Word documents from analysis results
    """
    
    def __init__(self):
        self.document = None
    
    def create_results_document(self, results_package: Dict) -> str:
        """
        Create a Word document with all analysis results
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            doc = Document()
            
            # Title
            title = doc.add_heading('Data Analysis Results', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add date
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
            doc.add_paragraph()
            
            # Methods section
            if 'methods' in results_package:
                doc.add_heading('Methods', level=1)
                doc.add_paragraph(results_package['methods'])
            
            # Results section
            doc.add_heading('Results', level=1)
            
            # Add each analysis result
            if 'analyses' in results_package:
                for i, analysis in enumerate(results_package['analyses'], 1):
                    doc.add_heading(f'{analysis.get("name", f"Analysis {i}")}', level=2)
                    
                    # Add table if present
                    if 'table' in analysis:
                        self._add_table(doc, analysis['table'])
                    
                    # Add narrative if present
                    if 'narrative' in analysis:
                        doc.add_paragraph(analysis['narrative'])
                    
                    doc.add_paragraph()
            
            # Save document
            output_path = f'results_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            doc.save(output_path)
            
            return output_path
            
        except ImportError:
            return None
    
    def _add_table(self, doc, table_data: Dict):
        """Add a table to the document"""
        from docx import Document
        from docx.shared import Pt
        
        if 'dataframe' in table_data and table_data['dataframe']:
            data = table_data['dataframe']
            
            if isinstance(data, list) and len(data) > 0:
                # Create table
                columns = list(data[0].keys())
                table = doc.add_table(rows=1, cols=len(columns))
                table.style = 'Table Grid'
                
                # Header row
                header_cells = table.rows[0].cells
                for i, col in enumerate(columns):
                    header_cells[i].text = str(col)
                
                # Data rows
                for row_data in data:
                    row_cells = table.add_row().cells
                    for i, col in enumerate(columns):
                        row_cells[i].text = str(row_data.get(col, ''))


class PythonCodeGenerator:
    """
    Generate Python code that reproduces the analysis
    """
    
    @staticmethod
    def generate_analysis_code(analyses: List[Dict]) -> str:
        """Generate Python code for reproducing analyses"""
        code_lines = [
            "# Data Analysis Code",
            "# Auto-generated by Students Companion Data Analysis Lab",
            f"# Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "import pandas as pd",
            "import numpy as np",
            "from scipy import stats",
            "import statsmodels.api as sm",
            "",
            "# Load your data",
            "# df = pd.read_csv('your_data.csv')",
            "",
        ]
        
        for analysis in analyses:
            code_lines.append(f"# --- {analysis.get('name', 'Analysis')} ---")
            
            analysis_type = analysis.get('type')
            
            if analysis_type == 'ttest':
                code_lines.extend([
                    f"# Independent Samples t-test",
                    f"group_var = '{analysis.get('group_var', 'group')}'",
                    f"outcome_var = '{analysis.get('outcome_var', 'outcome')}'",
                    "",
                    "groups = df[group_var].unique()",
                    "group1 = df[df[group_var] == groups[0]][outcome_var]",
                    "group2 = df[df[group_var] == groups[1]][outcome_var]",
                    "",
                    "# Run t-test",
                    "t_stat, p_value = stats.ttest_ind(group1, group2)",
                    "",
                    "# Cohen's d",
                    "pooled_std = np.sqrt(((len(group1)-1)*group1.var() + (len(group2)-1)*group2.var()) / (len(group1)+len(group2)-2))",
                    "cohens_d = (group1.mean() - group2.mean()) / pooled_std",
                    "",
                    "print(f't({len(group1)+len(group2)-2}) = {t_stat:.3f}, p = {p_value:.4f}')",
                    "print(f\"Cohen's d = {cohens_d:.3f}\")",
                    ""
                ])
            
            elif analysis_type == 'correlation':
                code_lines.extend([
                    f"# Correlation analysis",
                    f"var1 = '{analysis.get('var1', 'x')}'",
                    f"var2 = '{analysis.get('var2', 'y')}'",
                    "",
                    "r, p_value = stats.pearsonr(df[var1].dropna(), df[var2].dropna())",
                    "print(f'r = {r:.3f}, p = {p_value:.4f}')",
                    ""
                ])
            
            elif analysis_type == 'regression':
                code_lines.extend([
                    f"# Linear Regression",
                    f"outcome = '{analysis.get('outcome', 'y')}'",
                    f"predictors = {analysis.get('predictors', ['x'])}",
                    "",
                    "X = df[predictors]",
                    "y = df[outcome]",
                    "X = sm.add_constant(X)",
                    "",
                    "model = sm.OLS(y, X).fit()",
                    "print(model.summary())",
                    ""
                ])
            
            code_lines.append("")
        
        return "\n".join(code_lines)

